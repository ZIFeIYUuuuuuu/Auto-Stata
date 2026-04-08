"""Word 文档格式化工具

生成符合中文社科/理工类期刊排版规范的 Word 文档，
所有统计结果表格采用学术论文标准的三线表样式。
同时，为每张图自动插入图题、图注和专业文字解读。
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table

from agents.chart_commentary_agent import build_chart_commentary
from schemas import AnalysisStepResult


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = tc_pr.makeelement(qn("w:tcBorders"), {})
        tc_pr.append(borders)
    for edge, attrs in kwargs.items():
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = borders.makeelement(qn(f"w:{edge}"), {})
            borders.append(element)
        for attr_name, attr_val in attrs.items():
            element.set(qn(f"w:{attr_name}"), str(attr_val))


def apply_three_line_style(table: Table) -> None:
    thick = {"sz": "24", "val": "single", "color": "000000"}  # 1.5pt
    thin = {"sz": "12", "val": "single", "color": "000000"}   # 0.75pt
    none = {"sz": "0", "val": "none", "color": "000000"}

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if row_idx == 0:
                _set_cell_border(cell, top=thick, bottom=thin, left=none, right=none)
            elif row_idx == len(table.rows) - 1:
                _set_cell_border(cell, top=none, bottom=thick, left=none, right=none)
            else:
                _set_cell_border(cell, top=none, bottom=none, left=none, right=none)

            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Times New Roman"
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    r = run._element
                    r_fonts = r.find(qn("w:rPr"))
                    if r_fonts is not None:
                        fonts = r_fonts.find(qn("w:rFonts"))
                        if fonts is not None:
                            fonts.set(qn("w:eastAsia"), "宋体")

    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def _parse_stata_table(log_text: str) -> list[list[str]] | None:
    lines = log_text.strip().split("\n")
    table_lines: list[list[str]] = []
    in_table = False

    for line in lines:
        stripped = line.strip()
        if re.match(r"^[-+]+$", stripped):
            in_table = True
            continue
        if in_table and stripped:
            cleaned = re.sub(r"\|", " ", stripped)
            cells = cleaned.split()
            if cells:
                table_lines.append(cells)
        elif in_table and not stripped:
            if table_lines:
                break

    return table_lines if table_lines else None


def _add_heading(doc: Document, text: str, level: int = 2) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "黑体"
        r = run._element
        r_fonts = r.find(qn("w:rPr"))
        if r_fonts is not None:
            fonts = r_fonts.find(qn("w:rFonts"))
            if fonts is not None:
                fonts.set(qn("w:eastAsia"), "黑体")


def _add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        r = run._element
        r_fonts = r.find(qn("w:rPr"))
        if r_fonts is not None:
            fonts = r_fonts.find(qn("w:rFonts"))
            if fonts is not None:
                fonts.set(qn("w:eastAsia"), "宋体")


def _add_center_paragraph(doc: Document, text: str, font_name: str = "黑体", font_size: float = 10.5, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = font_name
    r = run._element
    r_fonts = r.find(qn("w:rPr"))
    if r_fonts is not None:
        fonts = r_fonts.find(qn("w:rFonts"))
        if fonts is not None:
            fonts.set(qn("w:eastAsia"), font_name)


def _add_figure_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.name = "黑体"
    r = run._element
    r_fonts = r.find(qn("w:rPr"))
    if r_fonts is not None:
        fonts = r_fonts.find(qn("w:rFonts"))
        if fonts is not None:
            fonts.set(qn("w:eastAsia"), "黑体")


def _add_figure_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"注：{text}")
    run.font.size = Pt(9)
    run.font.name = "宋体"
    r = run._element
    r_fonts = r.find(qn("w:rPr"))
    if r_fonts is not None:
        fonts = r_fonts.find(qn("w:rFonts"))
        if fonts is not None:
            fonts.set(qn("w:eastAsia"), "宋体")


def _add_table_from_log(doc: Document, log_text: str, caption: str) -> None:
    rows = _parse_stata_table(log_text)
    if not rows:
        _add_paragraph(doc, f"（{caption}：未能解析表格数据，原始日志见附录）")
        return

    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_p.add_run(caption)
    cap_run.font.size = Pt(10.5)
    cap_run.font.bold = True
    cap_run.font.name = "黑体"
    r = cap_run._element
    r_fonts = r.find(qn("w:rPr"))
    if r_fonts is not None:
        fonts = r_fonts.find(qn("w:rFonts"))
        if fonts is not None:
            fonts.set(qn("w:eastAsia"), "黑体")

    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < n_cols:
                cell = table.cell(i, j)
                cell.text = cell_text
                if i == 0:
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True

    apply_three_line_style(table)


def _add_raw_log_section(doc: Document, log_text: str, title: str) -> None:
    _add_heading(doc, title, level=3)
    p = doc.add_paragraph()
    run = p.add_run(log_text[:3000])
    run.font.size = Pt(8)
    run.font.name = "Courier New"


def _add_figure_with_commentary(
    doc: Document,
    img_path: str,
    step: str,
    log_text: str,
    fig_no: int,
) -> int:
    p = Path(img_path)
    if not p.exists():
        return fig_no

    commentary = build_chart_commentary(step=step, graph_path=img_path, log_text=log_text)

    # 标题中的图号替换
    title = commentary["title"].replace("图1", f"图{fig_no}").replace("图2", f"图{fig_no}").replace("图X", f"图{fig_no}")

    _add_figure_caption(doc, title)
    doc.add_picture(str(p), width=Cm(14))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_figure_note(doc, commentary["caption"])
    _add_paragraph(doc, commentary["commentary"])

    return fig_no + 1


STEP_TITLES = {
    "descriptive": ("一、描述性统计分析", "表1 描述性统计结果"),
    "vif": ("二、多重共线性检验（VIF）", "表2 VIF 检验结果"),
    "baseline_regression": ("三、基准回归结果", "表3 基准回归结果"),
    "panel": ("四、面板数据分析", "表4 面板回归结果与 Hausman 检验"),
    "robustness": ("五、稳健性检验", "表5 稳健性检验结果"),
    "heterogeneity": ("六、异质性分析", "表6 异质性分析（分组回归）"),
    "iv_2sls": ("七、内生性处理（2SLS）", "表7 工具变量回归结果"),
}


def build_report_docx(
    results: list[AnalysisStepResult],
    ai_report: str,
    output_path: Path,
) -> Path:
    doc = Document()

    title = doc.add_heading("AutoStata-Insight 实证分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fig_no = 1

    for r in results:
        if r.step not in STEP_TITLES:
            continue

        section_title, table_caption = STEP_TITLES[r.step]
        _add_heading(doc, section_title)

        if r.error:
            _add_paragraph(doc, f"本步骤执行出错：{r.error}")
            continue

        if r.log:
            _add_table_from_log(doc, r.log, table_caption)

        if r.graphs:
            for img_path in r.graphs:
                fig_no = _add_figure_with_commentary(
                    doc=doc,
                    img_path=img_path,
                    step=r.step,
                    log_text=r.log,
                    fig_no=fig_no,
                )

    _add_heading(doc, "八、AI 综合解读报告", level=1)

    for line in ai_report.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            _add_heading(doc, line.lstrip("# "), level=2)
        elif line.startswith("### "):
            _add_heading(doc, line.lstrip("# "), level=3)
        elif line.startswith("# "):
            _add_heading(doc, line.lstrip("# "), level=1)
        else:
            _add_paragraph(doc, line)

    _add_heading(doc, "附录：Stata 原始输出日志", level=1)
    for r in results:
        if r.log and r.step in STEP_TITLES:
            title_text = STEP_TITLES[r.step][0]
            _add_raw_log_section(doc, r.log, f"附录 - {title_text}")

    doc.save(str(output_path))
    return output_path